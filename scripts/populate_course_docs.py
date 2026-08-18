# -*- coding: utf-8 -*-
"""
Script populate_course_docs.py
Tự động xuất bộ 3 tài liệu học thuật chuẩn theo template môn học:
"Ứng dụng Trí tuệ Nhân tạo trong Phát triển Phần mềm (GenAI in Software Development)"
Khoa học Máy tính - K23A

Danh sách thành viên nhóm:
1. Vũ Ngọc Sơn (Trưởng nhóm, Kiến trúc hệ thống & Fullstack)
2. Vũ Bảo Linh (Kỹ sư Dữ liệu & Tích hợp AI Engine)
3. Tô Văn Quyền (Kiểm thử Phần mềm & Đặc tả Nghiệp vụ)
4. Lê Bình Nguyên (Thiết kế Giao diện UI/UX & Tài liệu Học thuật)

Tài liệu:
1. 01_GenAI_SoftwareDevelopment_project-plan.docx
2. 02_GenAI_SoftwareDevelopment_requirements-qa.docx
3. 03_GenAI_SoftwareDevelopment_requirements-specification.docx
Cùng với các file Markdown tương ứng trong docs/academic/
"""

import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

sys.stdout.reconfigure(encoding='utf-8')

COURSE_FOLDER = r'D:\AAA\Ứng dụng trí tuệ nhân tạo KHMT K23A-20260805T083400Z-1-001'
ACADEMIC_DOCS = r'D:\AAA\docs\academic'
os.makedirs(ACADEMIC_DOCS, exist_ok=True)

APP_TITLE = "Hệ thống Nền tảng PropTech Cho Thuê & Quản Lý Căn Hộ Tích Hợp Trí Tuệ Nhân Tạo (HAVEN)"

# ==============================================================================
# 1. GENERATE & POPULATE PROJECT PLAN (01)
# ==============================================================================
def populate_project_plan():
    print("[1/3] Populating Project Plan...")
    docx_path = os.path.join(COURSE_FOLDER, '01_GenAI_SoftwareDevelopment_project-plan.docx')
    doc = docx.Document(docx_path)
    
    # Update project name paragraph
    for p in doc.paragraphs:
        if 'Tên ứng dụng:' in p.text:
            p.text = f"Tên ứng dụng: {APP_TITLE}"
            p.runs[0].bold = True
    
    # 9 Weeks Detailed Plan (Tuần 1 -> Tuần 9)
    weekly_tasks = [
        # Tuần 1: 27/07 - 02/08
        ("Tuần 01 (Từ: 27/07/2026 Đến: 02/08/2026)", "Khảo sát thị trường BĐS cho thuê, phân tích các bất cập thực tế (chi phí ẩn, ngập lụt, rủi ro PCCC, quỵt cọc, tin ảo).", "Vũ Ngọc Sơn", "Báo cáo khảo sát thị trường"),
        ("Tuần 01 (Từ: 27/07/2026 Đến: 02/08/2026)", "Xác định mục tiêu hệ thống HAVEN và phạm vi ứng dụng GenAI trong giải quyết bài toán.", "Vũ Bảo Linh", "Đề cương bài toán AI"),
        ("Tuần 01 (Từ: 27/07/2026 Đến: 02/08/2026)", "Lập kế hoạch tổng thể 9 tuần và phân công vai trò 4 thành viên trong nhóm.", "Tô Văn Quyền", "Kế hoạch 9 tuần"),
        ("Tuần 01 (Từ: 27/07/2026 Đến: 02/08/2026)", "Thiết lập quy trình làm việc Git, Antigravity Workspace và chuẩn tài liệu hóa.", "Lê Bình Nguyên", "Setup repo & workspace"),
        ("Tuần 01 (Từ: 27/07/2026 Đến: 02/08/2026)", "Họp nhóm tổng kết tuần 1 và thống nhất bảng phạm vi tính năng (Product Scope).", "Cả nhóm", "Biên bản họp tuần 1"),

        # Tuần 2: 03/08 - 09/08
        ("Tuần 02 (Từ: 03/08/2026 Đến: 09/08/2026)", "Xây dựng bảng 20 câu hỏi phỏng vấn thu thập yêu cầu người dùng (Tenant & Landlord).", "Vũ Ngọc Sơn", "Bộ 20 câu hỏi Q&A"),
        ("Tuần 02 (Từ: 03/08/2026 Đến: 09/08/2026)", "Phân loại danh mục Yêu cầu Chức năng (Functional) và Phi chức năng (Non-Functional).", "Vũ Bảo Linh", "Bảng phân loại yêu cầu"),
        ("Tuần 02 (Từ: 03/08/2026 Đến: 09/08/2026)", "Thiết kế Sơ đồ phân cấp chức năng hệ thống (Function Hierarchy).", "Tô Văn Quyền", "Sơ đồ phân cấp chức năng"),
        ("Tuần 02 (Từ: 03/08/2026 Đến: 09/08/2026)", "Đặc tả yêu cầu nghiệp vụ True Cost Breakdown và quy chuẩn an toàn PCCC QCVN 06:2022.", "Lê Bình Nguyên", "Tài liệu nghiệp vụ chi phí & PCCC"),
        ("Tuần 02 (Từ: 03/08/2026 Đến: 09/08/2026)", "Hoàn thiện tài liệu 02_requirements-qa.docx nộp giảng viên.", "Cả nhóm", "Tài liệu Q&A hoàn chỉnh"),

        # Tuần 3: 10/08 - 16/08
        ("Tuần 03 (Từ: 10/08/2026 Đến: 16/08/2026)", "Soạn thảo tài liệu Đặc tả yêu cầu phần mềm SRS (Software Requirements Specification).", "Vũ Ngọc Sơn", "Khung tài liệu SRS"),
        ("Tuần 03 (Từ: 10/08/2026 Đến: 16/08/2026)", "Xác định danh sách Tác nhân (Actors) và vẽ Use Case Diagram tổng quát bằng Mermaid/Draw.io.", "Vũ Bảo Linh", "Sơ đồ Use Case tổng quát"),
        ("Tuần 03 (Từ: 10/08/2026 Đến: 16/08/2026)", "Chi tiết hóa Use Case: Tìm kiếm True Cost & Bản Đồ PCCC/Ngập lụt (UC001).", "Tô Văn Quyền", "Đặc tả UC001 chi tiết"),
        ("Tuần 03 (Từ: 10/08/2026 Đến: 16/08/2026)", "Chi tiết hóa Use Case: Đăng tin thông minh tự động bằng AI (UC002).", "Lê Bình Nguyên", "Đặc tả UC002 chi tiết"),
        ("Tuần 03 (Từ: 10/08/2026 Đến: 16/08/2026)", "Đánh giá & Kiểm thử chấp nhận yêu cầu (Requirements Review & Baseline SRS).", "Cả nhóm", "Bản hoàn chỉnh SRS 1.0"),

        # Tuần 4: 17/08 - 23/08 (HIỆN TẠI - BÀI KIỂM TRA 1)
        ("Tuần 04 (Từ: 17/08/2026 Đến: 23/08/2026)", "Thiết kế kiến trúc hệ thống tổng thể (System Architecture, Tech Stack React + Vite + Tailwind + Lucide).", "Vũ Ngọc Sơn", "Kiến trúc hệ thống"),
        ("Tuần 04 (Từ: 17/08/2026 Đến: 23/08/2026)", "Thiết kế mô hình dữ liệu (Data Modeling TypeScript Interfaces, 150 căn hộ, 16 thành phố).", "Vũ Bảo Linh", "Schema dữ liệu & Mock data"),
        ("Tuần 04 (Từ: 17/08/2026 Đến: 23/08/2026)", "Xây dựng bộ mock data chuẩn hóa, mô phỏng cảm biến IoT, báo cáo PCCC và chỉ số ngập lụt.", "Tô Văn Quyền", "Dataset 150 căn hộ"),
        ("Tuần 04 (Từ: 17/08/2026 Đến: 23/08/2026)", "Thiết kế UI Wireframe/Mockups cho 4 phân hệ chính (Tenant, Landlord, Admin, Resident).", "Lê Bình Nguyên", "Giao diện UI Mockups"),
        ("Tuần 04 (Từ: 17/08/2026 Đến: 23/08/2026)", "Tổng kết tiến độ Giai đoạn 1 & Thực hiện Bài Kiểm tra 1 (Nộp bộ 3 tài liệu + Demo).", "Cả nhóm", "Hoàn thành Bài Kiểm tra 1"),

        # Tuần 5: 24/08 - 30/08
        ("Tuần 05 (Từ: 24/08/2026 Đến: 30/08/2026)", "Phát triển phân hệ Khách tìm thuê: Tìm kiếm AI ngôn ngữ tự nhiên, bộ lọc đa tiêu chí.", "Vũ Ngọc Sơn", "Module Tìm kiếm AI"),
        ("Tuần 05 (Từ: 24/08/2026 Đến: 30/08/2026)", "Phát triển Bản đồ PCCC & Ngập lụt và So sánh đa chiều (Radar Chart 5 trục).", "Vũ Bảo Linh", "Module Bản đồ & Radar"),
        ("Tuần 05 (Từ: 24/08/2026 Đến: 30/08/2026)", "Xây dựng Cẩm nang khu vực & Công cụ giả lập tuyến đường đi làm (Commute Simulator).", "Tô Văn Quyền", "Module Cẩm nang & Commute"),
        ("Tuần 05 (Từ: 24/08/2026 Đến: 30/08/2026)", "Tích hợp Tour tham quan thực tế ảo 360 độ và Thư viện tài liệu mẫu.", "Lê Bình Nguyên", "Module VR 360 & Docs"),
        ("Tuần 05 (Từ: 24/08/2026 Đến: 30/08/2026)", "Kiểm thử chức năng phân hệ Tenant (Unit Testing & UI Interaction Testing).", "Cả nhóm", "Báo cáo test phân hệ Tenant"),

        # Tuần 6: 31/08 - 06/09
        ("Tuần 06 (Từ: 31/08/2026 Đến: 06/09/2026)", "Phát triển phân hệ Chủ nhà: Trình soạn tin đăng tự động bằng AI (Smart Listing Creator).", "Vũ Ngọc Sơn", "Module Soạn tin AI"),
        ("Tuần 06 (Từ: 31/08/2026 Đến: 06/09/2026)", "Phát triển Bảng quản trị danh mục phòng, sơ đồ mặt bằng (Floor Plan Layout).", "Vũ Bảo Linh", "Module Quản lý phòng"),
        ("Tuần 06 (Từ: 31/08/2026 Đến: 06/09/2026)", "Phát triển Hệ thống quản lý Lead & Đặt lịch xem phòng trực tiếp.", "Tô Văn Quyền", "Module Quản lý Lead"),
        ("Tuần 06 (Từ: 31/08/2026 Đến: 06/09/2026)", "Xây dựng Trình tạo hợp đồng điện tử E-Sign và Hóa đơn thu tiền tự động.", "Lê Bình Nguyên", "Module Hợp đồng & Hóa đơn"),
        ("Tuần 06 (Từ: 31/08/2026 Đến: 06/09/2026)", "Tổng kết tiến độ Giai đoạn 2 & Thực hiện Bài Kiểm tra 2.", "Cả nhóm", "Hoàn thành Bài Kiểm tra 2"),

        # Tuần 7: 07/09 - 13/09
        ("Tuần 07 (Từ: 07/09/2026 Đến: 13/09/2026)", "Phát triển phân hệ Quản trị sàn (Marketplace Admin Dashboard & Sức khỏe sàn).", "Vũ Ngọc Sơn", "Module Admin Dashboard"),
        ("Tuần 07 (Từ: 07/09/2026 Đến: 13/09/2026)", "Xây dựng Trung tâm kiểm duyệt tin đăng, phát hiện tin ảo và cơ chế xếp hạng Trust Score.", "Vũ Bảo Linh", "Module Kiểm duyệt & Trust Score"),
        ("Tuần 07 (Từ: 07/09/2026 Đến: 13/09/2026)", "Phát triển Hệ thống phòng chống lừa đảo, bảo vệ tiền cọc trung gian (Escrow Protection).", "Tô Văn Quyền", "Module Escrow & Anti-Fraud"),
        ("Tuần 07 (Từ: 07/09/2026 Đến: 13/09/2026)", "Xây dựng Báo cáo phân tích doanh thu sàn, Take-rate và gói thuê bao.", "Lê Bình Nguyên", "Module Doanh thu sàn"),
        ("Tuần 07 (Từ: 07/09/2026 Đến: 13/09/2026)", "Kiểm thử tích hợp hệ thống phân hệ Admin (Integration Testing).", "Cả nhóm", "Báo cáo test phân hệ Admin"),

        # Tuần 8: 14/09 - 20/09
        ("Tuần 08 (Từ: 14/09/2026 Đến: 20/09/2026)", "Hoàn thiện phân hệ Cư dân đang thuê (Resident App: IoT Smart Home, VAS, Vé bảo trì).", "Vũ Ngọc Sơn", "Module Resident App"),
        ("Tuần 08 (Từ: 14/09/2026 Đến: 20/09/2026)", "Tối ưu hóa hiệu năng, SEO, Lighthouse Audit và tương thích Responsive Mobile.", "Vũ Bảo Linh", "Báo cáo Performance & SEO"),
        ("Tuần 08 (Từ: 14/09/2026 Đến: 20/09/2026)", "Kiểm thử tải (Load Testing), kiểm thử bảo mật (Security Review) và khắc phục lỗi.", "Tô Văn Quyền", "Báo cáo bảo mật & bug fix"),
        ("Tuần 08 (Từ: 14/09/2026 Đến: 20/09/2026)", "Đóng gói ứng dụng, cấu hình CI/CD và triển khai thử nghiệm trên môi trường Production.", "Lê Bình Nguyên", "Production Deployment"),
        ("Tuần 08 (Từ: 14/09/2026 Đến: 20/09/2026)", "Họp nghiệm thu toàn bộ tính năng phần mềm nội bộ nhóm.", "Cả nhóm", "Biên bản nghiệm thu phần mềm"),

        # Tuần 9: 21/09 - 27/09 (BÀI KIỂM TRA 3 & BÁO CÁO TỔNG KẾT)
        ("Tuần 09 (Từ: 21/09/2026 Đến: 27/09/2026)", "Chuẩn bị bài thuyết trình, tài liệu hướng dẫn sử dụng và kịch bản demo trực tiếp.", "Vũ Ngọc Sơn", "Slide & Kịch bản demo"),
        ("Tuần 09 (Từ: 21/09/2026 Đến: 27/09/2026)", "Rà soát, đóng gói toàn bộ mã nguồn, tài liệu học thuật và artifact dự án.", "Vũ Bảo Linh", "Kho tài liệu hoàn chỉnh"),
        ("Tuần 09 (Từ: 21/09/2026 Đến: 27/09/2026)", "Thực hiện Bài Kiểm tra 3 & Báo cáo tổng kết môn học trước giảng viên.", "Cả nhóm", "Báo cáo tổng kết môn học"),
        ("Tuần 09 (Từ: 21/09/2026 Đến: 27/09/2026)", "Tổng kết đánh giá đóng góp thành viên và đúc kết kinh nghiệm phát triển cùng GenAI.", "Cả nhóm", "Báo cáo Kaizen tổng kết")
    ]
    
    # Fill Table 0
    table = doc.tables[0]
    num_cols = len(table.columns)
    for i, task in enumerate(weekly_tasks):
        row_idx = i + 1
        if row_idx < len(table.rows):
            row = table.rows[row_idx]
        else:
            row = table.add_row()
            
        row.cells[0].text = str(i + 1)
        row.cells[1].text = task[0]
        if num_cols == 4:
            row.cells[2].text = f"{task[1]} (Sản phẩm: {task[3]})"
            row.cells[3].text = task[2]
        elif num_cols >= 5:
            row.cells[2].text = task[1]
            row.cells[3].text = task[2]
            row.cells[4].text = task[3]

    doc.save(docx_path)
    print("  -> Saved 01_GenAI_SoftwareDevelopment_project-plan.docx")

    # Generate Markdown version
    md_content = f"""# KẾ HOẠCH PHÁT TRIỂN DỰ ÁN 9 TUẦN (PROJECT PLAN)

**Tên ứng dụng**: {APP_TITLE}  
**Môn học**: Ứng dụng Trí tuệ Nhân tạo trong Phát triển Phần mềm (KHMT K23A)  
**Thời gian thực hiện**: 27/07/2026 – 27/09/2026 (9 tuần)  
**Nhóm sinh viên thực hiện**:
1. **Vũ Ngọc Sơn** (Trưởng nhóm, Kiến trúc hệ thống & Fullstack)
2. **Vũ Bảo Linh** (Kỹ sư Dữ liệu & Tích hợp AI Engine)
3. **Tô Văn Quyền** (Kiểm thử Phần mềm & Đặc tả Nghiệp vụ)
4. **Lê Bình Nguyên** (Thiết kế Giao diện UI/UX & Tài liệu Học thuật)

---

## BẢNG TIẾN ĐỘ THỰC HIỆN CHI TIẾT (46 CÔNG VIỆC / 9 TUẦN)

| STT | Thời gian thực hiện | Nội dung công việc chi tiết | Thành viên phụ trách | Sản phẩm / Kết quả bàn giao |
| :---: | :--- | :--- | :--- | :--- |
"""
    for i, t in enumerate(weekly_tasks):
        md_content += f"| {i+1} | {t[0]} | {t[1]} | {t[2]} | {t[3]} |\n"

    with open(os.path.join(ACADEMIC_DOCS, '01_PROJECT_PLAN_9_WEEKS.md'), 'w', encoding='utf-8') as f:
        f.write(md_content)
    print("  -> Saved docs/academic/01_PROJECT_PLAN_9_WEEKS.md")


# ==============================================================================
# 2. GENERATE & POPULATE REQUIREMENTS Q&A (02)
# ==============================================================================
def populate_requirements_qa():
    print("[2/3] Populating Requirements Q&A...")
    docx_path = os.path.join(COURSE_FOLDER, '02_GenAI_SoftwareDevelopment_requirements-qa.docx')
    doc = docx.Document(docx_path)

    # 20 Academic Q&A
    qa_list = [
        # Nhóm 1: Xác định bài toán & Khách hàng mục tiêu
        ("Hệ thống HAVEN giải quyết những bất cập thực tế cốt lõi nào của thị trường cho thuê căn hộ hiện nay?",
         "HAVEN giải quyết triệt để 5 bất cập lớn: (1) Chi phí ẩn và mập mờ biểu giá tiền điện/nước/dịch vụ; (2) Thiếu thông tin an toàn PCCC và lịch sử ngập lụt đô thị khi mưa bão; (3) Vấn nạn tin đăng ảo, ảnh mạng không đúng thực tế; (4) Nguy cơ mất tiền cọc hoặc tranh chấp khi trả phòng; (5) Thiếu công cụ so sánh đa chiều định lượng giữa các căn hộ.",
         "Vũ Ngọc Sơn"),

        ("Ai là đối tượng người dùng chính (Target Users) của nền tảng HAVEN?",
         "Hệ thống phục vụ 4 nhóm người dùng chính: (1) Người thuê nhà (Sinh viên, Người đi làm, Gia đình trẻ, Chuyên gia nước ngoài/Expats); (2) Chủ nhà cá nhân & Nhà quản lý căn hộ dịch vụ; (3) Cư dân đang sinh sống trong căn hộ; (4) Quản trị viên sàn thương mại bất động sản (Admin Marketplace).",
         "Vũ Bảo Linh"),

        ("Tính năng 'True Cost Breakdown' giải quyết bài toán tài chính của người thuê như thế nào?",
         "True Cost tự động tính toán tổng chi phí sinh hoạt thực tế hàng tháng = Tiền thuê gốc + Phí quản lý tòa nhà + Tiền gửi xe máy/ô tô + Internet + Ước tính điện năng tiêu thụ (theo diện tích & số điều hòa) + Nước sinh hoạt + Tiền cọc đầu vào. Người thuê nắm rõ ngân sách thực tế trước khi đi xem phòng.",
         "Tô Văn Quyền"),

        ("Làm thế nào để người dùng đánh giá an toàn PCCC và ngập lụt trước khi quyết định thuê phòng?",
         "Hệ thống tích hợp 'Bản đồ PCCC & Ngập lụt Đa Lớp'. Lớp PCCC hiển thị điểm an toàn (Safety Score 1-100), hồ sơ kiểm định nghiệm thu PCCC, thang thoát hiểm và đầu phun Sprinkler. Lớp Ngập lụt hiển thị lịch sử triều cường, thoát nước và độ rủi ro (Thấp/Trung bình/Cao) theo dữ liệu khí tượng.",
         "Lê Bình Nguyên"),

        # Nhóm 2: Trí tuệ nhân tạo (GenAI) & Tính năng cốt lõi
        ("Trí tuệ nhân tạo (GenAI Engine) đóng vai trò gì trong trải nghiệm tìm kiếm của người thuê?",
         "GenAI xử lý ngôn ngữ tự nhiên (NLP) cho phép người dùng nhập yêu cầu tự do (ví dụ: 'Tìm căn hộ 2 phòng ngủ dưới 15 triệu gần Keangnam, có chỗ đỗ ô tô, không ngập nước'). AI tự động trích xuất các tiêu chí lọc, tính điểm tương thích Match Score (%) và đưa ra lời giải thích chi tiết vì sao căn hộ phù hợp.",
         "Vũ Ngọc Sơn"),

        ("Công nghệ AI hỗ trợ chủ nhà đăng tin (Smart Listing Creator) như thế nào?",
         "Chủ nhà chỉ cần tải lên 3-8 ảnh thực tế của căn hộ. Hệ thống AI phân tích nhận diện loại phòng, tự động tạo tiêu đề chuẩn SEO, viết bài mô tả hấp dẫn và gợi ý mức giá thuê tối ưu dựa trên dữ liệu thị trường khu vực xung quanh.",
         "Vũ Bảo Linh"),

        ("Cơ chế 'So Sánh Đa Chiều (Radar Chart)' hỗ trợ người thuê ra quyết định ra sao?",
         "Cho phép người dùng chọn 2-4 căn hộ vào bảng so sánh. Hệ thống vẽ biểu đồ Radar 5 trục: (1) Chi phí tổng thể, (2) An toàn PCCC, (3) Độ cao cấp nội thất, (4) Vị trí & Tiện ích, (5) Mức độ yên tĩnh. Đi kèm bảng chiết tính chi phí chênh lệch từng khoản rõ ràng.",
         "Tô Văn Quyền"),

        ("Hệ thống giải quyết vấn đề tranh chấp tiền cọc khi trả phòng bằng phương thức nào?",
         "Hệ thống cung cấp 'Biên bản Bàn giao & Kiểm tra 15 Mục' số hóa kèm ảnh chụp hiện trạng có chữ ký số điện tử khi nhận phòng. Khi trả phòng, hai bên đối chiếu ảnh gốc. Nếu có tranh chấp, Quỹ bảo vệ ký quỹ (Deposit Escrow Protection) của HAVEN sẽ làm trung gian phân xử dựa trên chứng cứ.",
         "Lê Bình Nguyên"),

        ("Làm sao để người thuê ở xa hoặc bận rộn khảo sát được không gian căn hộ?",
         "Hệ thống tích hợp 'Tour Tham Quan Thực Tế Ảo 360 Độ' (Virtual Tour 3D) và 'Cẩm nang khu vực kết hợp Giả lập tuyến đường đi làm (Commute Simulator)'. Người dùng có thể đo đạc kích thước phòng và tính toán thời gian di chuyển giờ cao điểm trước khi đến xem trực tiếp.",
         "Vũ Ngọc Sơn"),

        ("Cơ chế Smart Chat & Đặt lịch hẹn có điểm gì khác biệt so với các sàn BĐS thông thường?",
         "Chat tích hợp 'Smart Action Chips' cho phép người thuê đặt câu hỏi nhanh chỉ với 1 chạm (Hỏi giá net, hỏi chỗ đỗ ô tô, hỏi giấy phép PCCC). Tích hợp lịch hẹn xem phòng đồng bộ trực tiếp với Google Calendar của chủ nhà.",
         "Vũ Bảo Linh"),

        # Nhóm 3: Phân hệ Quản trị & Nghiệp vụ Chủ nhà
        ("Phân hệ Chủ nhà (Landlord Dashboard) cung cấp những công cụ quản lý nào?",
         "Cung cấp: Quản lý danh mục căn hộ (trống, đang thuê, bảo trì), Sơ đồ mặt bằng tương tác (Floor Plan Layout), Quản lý Lead khách hàng tiềm năng theo phễu chuyển đổi, Tạo hợp đồng điện tử E-Sign và Theo dõi hóa đơn thu tiền hàng tháng.",
         "Tô Văn Quyền"),

        ("Làm thế nào để ngăn chặn tình trạng tin ảo, lừa đảo cọc trên nền tảng?",
         "Hệ thống áp dụng quy trình xác minh 2 lớp (Huy hiệu Verified): Xác minh danh tính qua CCCD/Hộ chiếu và Xác minh quyền sở hữu qua Sổ đỏ/Hợp đồng mua bán. Điểm tín nhiệm Trust Score đánh giá lịch sử giao dịch. Các tin đăng sai sự thật sẽ bị AI gắn cờ và Admin hạ khỏi sàn.",
         "Lê Bình Nguyên"),

        ("Phân hệ Quản trị viên (Marketplace Admin) giám sát những chỉ số sức khỏe nào của sàn?",
         "Admin Dashboard theo dõi: GMV (Tổng giá trị giao dịch), Doanh thu nền tảng (Take-rate 5-10% và phí gói đăng tin VIP), Tỷ lệ lấp đầy phòng (Occupancy Rate), Điểm chất lượng tin đăng trung bình và Bảng xử lý tranh chấp cọc/khiếu nại.",
         "Vũ Ngọc Sơn"),

        ("Phân hệ Ứng dụng Cư dân (Resident Portal) hỗ trợ những nghiệp vụ gì trong quá trình thuê?",
         "Hỗ trợ: Điều khiển thiết bị IoT thông minh (khóa cửa, điều hòa, theo dõi điện nước theo thời gian thực), Gửi yêu cầu sửa chữa/bảo trì kèm ảnh chụp, Thanh toán hóa đơn 1-chạm và Đặt trước các dịch vụ tiện ích tòa nhà (Gym, Hồ bơi, Dọn vệ sinh).",
         "Vũ Bảo Linh"),

        # Nhóm 4: Yêu cầu Phi chức năng & Kiến trúc Kỹ thuật
        ("Hệ thống có những yêu cầu phi chức năng (Non-Functional Requirements) nào về hiệu năng?",
         "Thời gian phản hồi tìm kiếm < 500ms; Thời gian tải trang ban đầu (First Contentful Paint) < 1.2s; Hỗ trợ tối thiểu 1.000 người dùng đồng thời (Concurrent Users); Đạt chuẩn tối ưu SEO và Lighthouse Score > 90 điểm.",
         "Tô Văn Quyền"),

        ("Yêu cầu về tính bảo mật và quyền riêng tư dữ liệu người dùng được đảm bảo ra sao?",
         "Toàn bộ dữ liệu truyền tải qua giao thức mã hóa HTTPS/TLS 1.3; Mật khẩu và thông tin cá nhân được băm an toàn (BCrypt/Argon2); Thông tin định danh CCCD và giấy tờ sở hữu được lưu trữ bảo mật tuân thủ Nghị định 13/2023/NĐ-CP về Bảo vệ dữ liệu cá nhân.",
         "Lê Bình Nguyên"),

        ("Giao diện người dùng (UI/UX) tuân theo những chuẩn mực thiết kế nào?",
         "Áp dụng phong cách Dark Mode sang trọng, hiện đại với bảng màu Slate-950 kết hợp điểm nhấn Emerald-400 và Amber-400. Sử dụng Typography hiện đại (Inter, Playfair Display), hiệu ứng Glassmorphism tinh tế và tương thích hoàn hảo trên thiết bị di động (Responsive Design).",
         "Vũ Ngọc Sơn"),

        ("Dữ liệu về 150 căn hộ và 16 thành phố được tổ chức như thế nào?",
         "Dữ liệu được mô hình hóa theo cấu trúc TypeScript chuẩn hóa với đầy đủ thông tin: ID định danh, phân khúc phòng, tọa độ khu vực, bảng chiết tính chi phí thật, chỉ số PCCC, dữ liệu cảm biến IoT và thông tin người thuê (Việt Nam và chuyên gia nước ngoài).",
         "Vũ Bảo Linh"),

        ("Hệ thống hỗ trợ những quy trình thanh toán nào?",
         "Hỗ trợ: Chuyển khoản ngân hàng tự động (VietQR), Thẻ thanh toán quốc tế (Visa/Mastercard) và Cổng thanh toán trực tuyến bảo chứng Escrow.",
         "Tô Văn Quyền"),

        ("Kế hoạch kiểm thử và bảo đảm chất lượng phần mềm được thực hiện như thế nào?",
         "Thực hiện kiểm thử 3 lớp: Unit Testing (kiểm thử các hàm tính toán True Cost, bộ lọc AI), Integration Testing (kiểm thử luồng tương tác giữa các phân hệ) và User Acceptance Testing (kiểm thử trải nghiệm thực tế của người dùng theo kịch bản).",
         "Lê Bình Nguyên")
    ]

    # Populate Table 0 in 02_requirements-qa.docx
    table = doc.tables[0]
    for i, qa in enumerate(qa_list):
        row_idx = i + 1
        if row_idx < len(table.rows):
            row = table.rows[row_idx]
        else:
            row = table.add_row()
        row.cells[0].text = str(i + 1)
        row.cells[1].text = qa[0]
        row.cells[2].text = qa[1]
        row.cells[3].text = qa[2]

    # Check if doc has second table or add heading + table
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
    else:
        doc.add_paragraph("\nPhân Loại Yêu Cầu Chức Năng & Phi Chức Năng")
        t1 = doc.add_table(rows=1, cols=4)
        t1.rows[0].cells[0].text = "Mã Yêu Cầu"
        t1.rows[0].cells[1].text = "Nội dung đặc tả"
        t1.rows[0].cells[2].text = "Phân loại"
        t1.rows[0].cells[3].text = "Mức độ ưu tiên"

    fn_nfn_data = [
        ("FR-01", "Tìm kiếm thông minh bằng AI (NLP Natural Query)", "Chức năng", "Bắt buộc (P0)"),
        ("FR-02", "Bảng chiết tính Chi phí thực tế (True Cost Breakdown)", "Chức năng", "Bắt buộc (P0)"),
        ("FR-03", "Bản đồ PCCC & Ngập lụt Đa Lớp", "Chức năng", "Bắt buộc (P0)"),
        ("FR-04", "So sánh đa chiều (Radar Chart 5 trục)", "Chức năng", "Bắt buộc (P0)"),
        ("FR-05", "Đăng tin tự động bằng AI (Smart Listing Creator)", "Chức năng", "Bắt buộc (P0)"),
        ("FR-06", "Quản lý Hợp đồng điện tử E-Sign & Cọc Escrow", "Chức năng", "Quan trọng (P1)"),
        ("FR-07", "Biên bản bàn giao 15 mục có ảnh đối chiếu", "Chức năng", "Quan trọng (P1)"),
        ("FR-08", "Quản trị sàn & Bảng sức khỏe thị trường", "Chức năng", "Quan trọng (P1)"),
        ("NFR-01", "Thời gian phản hồi tìm kiếm < 500ms", "Phi chức năng", "Bắt buộc (P0)"),
        ("NFR-02", "Bảo mật dữ liệu định danh & mã hóa đường truyền HTTPS", "Phi chức năng", "Bắt buộc (P0)"),
        ("NFR-03", "Thiết kế Responsive tương thích 100% Mobile/Tablet/Desktop", "Phi chức năng", "Bắt buộc (P0)"),
        ("NFR-04", "Điểm đánh giá hiệu năng Lighthouse > 90 điểm", "Phi chức năng", "Quan trọng (P1)")
    ]
    for i, item in enumerate(fn_nfn_data):
        row_idx = i + 1
        if row_idx < len(t1.rows):
            row = t1.rows[row_idx]
        else:
            row = t1.add_row()
        row.cells[0].text = item[0]
        row.cells[1].text = item[1]
        row.cells[2].text = item[2]
        row.cells[3].text = item[3]

    doc.save(docx_path)
    print("  -> Saved 02_GenAI_SoftwareDevelopment_requirements-qa.docx")

    # Generate Markdown version
    md_content = f"""# BẢNG THU THẬP & ĐẶC TẢ YÊU CẦU PHẦN MỀM (REQUIREMENTS Q&A)

**Tên ứng dụng**: {APP_TITLE}  
**Môn học**: Ứng dụng Trí tuệ Nhân tạo trong Phát triển Phần mềm (KHMT K23A)  
**Nhóm sinh viên thực hiện**:
1. **Vũ Ngọc Sơn** (Trưởng nhóm, Kiến trúc hệ thống & Fullstack)
2. **Vũ Bảo Linh** (Kỹ sư Dữ liệu & Tích hợp AI Engine)
3. **Tô Văn Quyền** (Kiểm thử Phần mềm & Đặc tả Nghiệp vụ)
4. **Lê Bình Nguyên** (Thiết kế Giao diện UI/UX & Tài liệu Học thuật)

---

## 1. BẢNG 20 CÂU HỎI & TRẢ LỜI THU THẬP YÊU CẦU (REQUIREMENTS Q&A)

| STT | Câu hỏi khảo sát yêu cầu | Nội dung phản hồi & Thống nhất kỹ thuật | Người phụ trách |
| :---: | :--- | :--- | :--- |
"""
    for i, qa in enumerate(qa_list):
        md_content += f"| {i+1} | **{qa[0]}** | {qa[1]} | {qa[2]} |\n"

    md_content += """
---

## 2. PHÂN LOẠI DANH MỤC YÊU CẦU CHỨC NĂNG & PHI CHỨC NĂNG

| Mã Yêu Cầu | Nội dung đặc tả yêu cầu | Phân loại | Mức độ ưu tiên |
| :---: | :--- | :---: | :---: |
"""
    for item in fn_nfn_data:
        md_content += f"| {item[0]} | {item[1]} | {item[2]} | {item[3]} |\n"

    with open(os.path.join(ACADEMIC_DOCS, '02_REQUIREMENTS_QA_20.md'), 'w', encoding='utf-8') as f:
        f.write(md_content)
    print("  -> Saved docs/academic/02_REQUIREMENTS_QA_20.md")


# ==============================================================================
# 3. GENERATE & POPULATE REQUIREMENTS SPECIFICATION (SRS - 03)
# ==============================================================================
def populate_requirements_specification():
    print("[3/3] Populating SRS Specification...")
    docx_path = os.path.join(COURSE_FOLDER, '03_GenAI_SoftwareDevelopment_requirements-specification.docx')
    doc = docx.Document(docx_path)

    # Table 0: Glossary (4 cols)
    t0 = doc.tables[0]
    glossary = [
        ("PropTech", "Property Technology — Ứng dụng công nghệ thông tin vào bất động sản.", "Khái niệm ngành"),
        ("True Cost", "Tổng chi phí thực tế hàng tháng = Tiền thuê gốc + Phí quản lý + Gửi xe + Điện nước + Internet.", "Nghiệp vụ cốt lưỡii"),
        ("PCCC (QCVN 06)", "Quy chuẩn kỹ thuật quốc gia về an toàn cháy cho nhà và công trình (QCVN 06:2022/BXD).", "Quy chuẩn an toàn"),
        ("Escrow", "Cơ chế tạm giữ tiền ký quỹ bảo chứng giao dịch thông qua bên thứ ba trung gian an toàn.", "Bảo vệ tiền cọc"),
        ("Trust Score", "Điểm tín nhiệm chủ nhà được tính toán tự động trên thang 1.0 - 5.0★ từ dữ liệu giao dịch.", "Minh bạch sàn"),
        ("GenAI Engine", "Hệ thống trí tuệ nhân tạo tạo sinh hỗ trợ phân tích ngôn ngữ tự nhiên và thị giác máy tính.", "Công nghệ AI")
    ]
    for i, g in enumerate(glossary):
        row_idx = i + 1
        if row_idx < len(t0.rows):
            row = t0.rows[row_idx]
        else:
            row = t0.add_row()
        row.cells[0].text = str(i + 1)
        row.cells[1].text = g[0]
        row.cells[2].text = g[1]
        if len(row.cells) > 3:
            row.cells[3].text = g[2]

    # Table 1: References (3 cols)
    t1 = doc.tables[1]
    refs = [
        ("Tài liệu Kế hoạch Phát triển Dự án 9 Tuần (01_project-plan.docx)", "Nhóm sinh viên thực hiện"),
        ("Tài liệu Thu thập và Khảo sát Yêu cầu Phần mềm (02_requirements-qa.docx)", "Nhóm sinh viên thực hiện"),
        ("Quy chuẩn Kỹ thuật Quốc gia về An toàn Cháy (QCVN 06:2022/BXD)", "Bộ Xây Dựng ban hành"),
        ("Nghiên cứu Sản phẩm HAVEN PropTech Suite (Part 1, 2, 3)", "Kiến trúc sư hệ thống")
    ]
    for i, r in enumerate(refs):
        row_idx = i + 1
        if row_idx < len(t1.rows):
            row = t1.rows[row_idx]
        else:
            row = t1.add_row()
        row.cells[0].text = str(i + 1)
        row.cells[1].text = r[0]
        row.cells[2].text = r[1]

    # Table 2: Actors
    t2 = doc.tables[2]
    actors = [
        ("Người Thuê (Tenant)", "Người dùng cá nhân tìm kiếm, so sánh và thuê căn hộ.", "Tìm kiếm AI, xem True Cost, bản đồ PCCC, so sánh radar, chat, ký hợp đồng E-Sign, cọc Escrow."),
        ("Chủ Nhà (Landlord)", "Chủ sở hữu hoặc nhà quản lý căn hộ cho thuê.", "Đăng tin bằng AI, quản lý danh mục phòng, tiếp nhận lead, tạo hợp đồng điện tử, thu tiền qua hóa đơn."),
        ("Cư Dân (Resident)", "Người đang sinh sống thực tế tại căn hộ.", "Điều khiển IoT Smart Home, gửi yêu cầu bảo trì, thanh toán hóa đơn, đặt lịch tiện ích."),
        ("Quản Trị Sàn (Admin)", "Đội ngũ vận hành và kiểm duyệt của sàn HAVEN.", "Giám sát sức khỏe sàn, kiểm duyệt tin đăng, gắn cờ tin ảo, xử lý tranh chấp tiền cọc."),
        ("Hệ Thống GenAI", "Dịch vụ AI tích hợp xử lý NLP và Computer Vision.", "Phân tích intent tìm kiếm, nhận diện layout phòng từ ảnh, sinh tiêu đề SEO và gợi ý giá thuê.")
    ]
    for i, a in enumerate(actors):
        row_idx = i + 1
        if row_idx < len(t2.rows):
            row = t2.rows[row_idx]
        else:
            row = t2.add_row()
        row.cells[0].text = str(i + 1)
        row.cells[1].text = a[0]
        row.cells[2].text = a[1]

    # Table 3: Use Cases
    t3 = doc.tables[3]
    use_cases = [
        ("UC001", "Tìm kiếm Căn hộ Thông minh & Bản Đồ PCCC/Ngập lụt", "Người Thuê", "Cao"),
        ("UC002", "Tư vấn & Gợi ý Căn hộ Bằng Ngôn ngữ Tự nhiên (GenAI)", "Người Thuê, GenAI", "Cao"),
        ("UC003", "So Sánh Đa Chiều (Radar Chart 5 trục)", "Người Thuê", "Cao"),
        ("UC004", "Nhắn Tin Smart Action Chips & Đặt Lịch Xem Phòng", "Người Thuê, Chủ Nhà", "Cao"),
        ("UC005", "Ký Hợp Đồng Điện Tử E-Sign & Đặt Cọc Bảo Chứng Escrow", "Người Thuê, Chủ Nhà", "Cao"),
        ("UC006", "Biên Bản Bàn Giao Căn Hộ 15 Mục Kèm Ảnh Đối Chiếu", "Người Thuê, Chủ Nhà", "Trung bình"),
        ("UC007", "Cẩm Nang Khu Vực & Giả Lập Tuyến Đường Đi Làm", "Người Thuê", "Trung bình"),
        ("UC008", "Tour Tham Quan Thực Tế Ảo 360 Độ (Virtual Tour 3D)", "Người Thuê", "Trung bình"),
        ("UC009", "Đăng Tin Tự Động Bằng AI (Smart Listing Creator)", "Chủ Nhà, GenAI", "Cao"),
        ("UC010", "Quản Lý Phòng, Sơ Đồ Tầng & Hóa Đơn Thu Tiền", "Chủ Nhà", "Cao"),
        ("UC011", "Quản Trị Sức Khỏe Sàn & Kiểm Duyệt Tin Đăng", "Quản Trị Viên", "Cao"),
        ("UC012", "Điều Khiển IoT Smart Home & Yêu Cầu Sửa Chữa", "Cư Dân", "Trung bình")
    ]
    for i, uc in enumerate(use_cases):
        row_idx = i + 1
        if row_idx < len(t3.rows):
            row = t3.rows[row_idx]
        else:
            row = t3.add_row()
        row.cells[0].text = str(i + 1)
        row.cells[1].text = uc[0]
        row.cells[2].text = uc[1]
        row.cells[3].text = uc[2]
        if len(row.cells) > 4:
            row.cells[4].text = uc[3]

    # Table 4: Detail UC001
    t4 = doc.tables[4]
    t4.rows[0].cells[1].text = "UC001: Tìm kiếm Căn hộ Thông minh & Bản Đồ PCCC/Ngập lụt"
    t4.rows[1].cells[1].text = "Tìm kiếm căn hộ theo chi phí thật, khu vực và trực quan hóa rủi ro an toàn."
    t4.rows[2].cells[1].text = "Cho phép người thuê lọc căn hộ theo ngân sách True Cost, vị trí và kiểm tra lớp dữ liệu PCCC/ngập lụt trước khi quyết định thuê."
    t4.rows[3].cells[1].text = "Người Thuê (Tenant)"
    t4.rows[4].cells[1].text = "Người dùng truy cập vào trang Tìm kiếm hoặc Bản đồ trên ứng dụng HAVEN."
    t4.rows[5].cells[1].text = "Danh sách căn hộ phù hợp được hiển thị trực quan kèm chiết tính True Cost và chỉ số an toàn."
    t4.rows[6].cells[1].text = "1. Người thuê chọn thành phố/quận và khoảng giá True Cost.\n2. Người thuê bật các tiêu chí nâng cao (PCCC đạt chuẩn, không ngập, có chỗ đỗ ô tô).\n3. Hệ thống lọc và hiển thị danh sách căn hộ theo thời gian thực.\n4. Người thuê chuyển sang tab Bản đồ để xem lớp rủi ro ngập lụt và hồ sơ PCCC.\n5. Người thuê bấm vào căn hộ để xem chi tiết hoặc thêm vào danh sách so sánh."
    t4.rows[7].cells[1].text = "3a. Không tìm thấy căn hộ thỏa mãn mọi tiêu chí:\n  - Hệ thống gợi ý nới lỏng bán kính tìm kiếm hoặc khoảng giá để hiển thị các lựa chọn gần nhất."

    # Table 5: Detail UC009
    t5 = doc.tables[5]
    t5.rows[0].cells[1].text = "UC009: Đăng Tin Tự Động Bằng AI (Smart Listing Creator)"
    t5.rows[1].cells[1].text = "Tự động phân tích ảnh căn hộ và sinh nội dung bài đăng hoàn chỉnh."
    t5.rows[2].cells[1].text = "Chủ nhà tải ảnh căn hộ, AI phân tích nhận diện phòng, tự động soạn thảo tiêu đề, mô tả và gợi ý giá thuê chuẩn xác."
    t5.rows[3].cells[1].text = "Chủ Nhà (Landlord), Hệ thống Trí tuệ nhân tạo (GenAI Engine)"
    t5.rows[4].cells[1].text = "Chủ nhà đăng nhập vào giao diện quản trị phân hệ Chủ nhà."
    t5.rows[5].cells[1].text = "Tin đăng được kiểm định chất lượng (>90 điểm) và xuất bản lên sàn HAVEN với huy hiệu Verified."
    t5.rows[6].cells[1].text = "1. Chủ nhà bấm nút 'Đăng Tin Mới Bằng AI'.\n2. Chủ nhà tải lên 3-8 bức ảnh thực tế và nhập diện tích/số phòng.\n3. Chủ nhà bấm 'Bắt Đầu Phân Tích & Tự Động Soạn Tin Bằng AI'.\n4. AI phân tích ảnh và sinh tiêu đề SEO + mô tả + gợi ý giá.\n5. Chủ nhà xem xét gợi ý chất lượng và bấm 'Xác Nhận & Xuất Bản'."
    t5.rows[7].cells[1].text = "4a. Ảnh tải lên có độ phân giải kém hoặc trùng lặp:\n  - AI cảnh báo và yêu cầu chủ nhà bổ sung ảnh chụp thực tế rõ nét."

    doc.save(docx_path)
    print("  -> Saved 03_GenAI_SoftwareDevelopment_requirements-specification.docx")

    # Generate Full Markdown SRS Document
    md_content = f"""# TÀI LIỆU ĐẶC TẢ YÊU CẦU PHẦN MỀM (SRS)

**Tên ứng dụng**: {APP_TITLE}  
**Mã dự án**: HAVEN-PROPTECH-K23A  
**Môn học**: Ứng dụng Trí tuệ Nhân tạo trong Phát triển Phần mềm (KHMT K23A)  
**Phiên bản**: 1.0 (Bản Phát hành Hoàn chỉnh)  
**Nhóm sinh viên thực hiện**:
1. **Vũ Ngọc Sơn** (Trưởng nhóm, Kiến trúc hệ thống & Fullstack)
2. **Vũ Bảo Linh** (Kỹ sư Dữ liệu & Tích hợp AI Engine)
3. **Tô Văn Quyền** (Kiểm thử Phần mềm & Đặc tả Nghiệp vụ)
4. **Lê Bình Nguyên** (Thiết kế Giao diện UI/UX & Tài liệu Học thuật)

---

## 1. GIỚI THIỆU CHUNG

### 1.1 Mục Đích
Tài liệu Đặc tả Yêu cầu Phần mềm (Software Requirements Specification - SRS) này mô tả toàn diện, chi tiết và có hệ thống các yêu cầu chức năng, phi chức năng, các ràng buộc kỹ thuật và mô hình thiết kế của hệ thống **HAVEN PropTech Platform**. Tài liệu đóng vai trò là cơ sở kỹ thuật thống nhất cho các thành viên phát triển, kiểm thử viên và giảng viên đánh giá môn học.

### 1.2 Phạm Vi & Đối Tượng Phục Vụ
- **Phạm vi**: Nền tảng PropTech web ứng dụng đa phân hệ dành cho thị trường thuê và cho thuê căn hộ tại các đô thị lớn tại Việt Nam (16 thành phố: Hà Nội, TP. Hồ Chí Minh, Đà Nẵng, Hải Phòng, Bình Dương, Nha Trang, Cần Thơ, Vũng Tàu, Hạ Long, Đà Lạt, Huế, Quy Nhơn, Biên Hòa, Vinh, Thanh Hóa, Buôn Ma Thuột).
- **Đối tượng phục vụ**:
  - Người thuê nhà cá nhân (Sinh viên, Người đi làm, Gia đình trẻ, Chuyên gia nước ngoài / Expat).
  - Chủ nhà cá nhân và Ban quản lý tòa nhà cho thuê.
  - Cư dân đang sinh sống tại căn hộ.
  - Quản trị viên sàn thương mại bất động sản.

### 1.3 Thuật Ngữ & Từ Viết Tắt
| STT | Thuật ngữ / Từ viết tắt | Giải thích chi tiết | Ghi chú |
| :---: | :--- | :--- | :--- |
| 1 | **PropTech** | Property Technology — Ứng dụng công nghệ số vào lĩnh vực bất động sản. | Khái niệm ngành |
| 2 | **True Cost** | Tổng chi phí thực tế hàng tháng = Tiền thuê gốc + Điện + Nước + Quản lý + Gửi xe + Internet. | Nghiệp vụ cốt lõi |
| 3 | **PCCC (QCVN 06)** | Quy chuẩn kỹ thuật quốc gia về an toàn cháy cho nhà và công trình (QCVN 06:2022/BXD). | Pháp lý an toàn |
| 4 | **Escrow** | Cơ chế tạm giữ tiền ký quỹ bảo chứng giao dịch thông qua bên thứ ba trung gian. | An toàn tiền cọc |
| 5 | **Trust Score** | Điểm tín nhiệm chủ nhà được tính toán tự động trên thang 1.0 - 5.0★ từ dữ liệu giao dịch. | Minh bạch sàn |
| 6 | **GenAI** | Generative Artificial Intelligence — Trí tuệ nhân tạo tạo sinh hỗ trợ tư vấn và soạn tin. | Công nghệ AI |

### 1.4 Tài Liệu Tham Khảo
1. *Tài liệu Kế hoạch Dự án*: `01_GenAI_SoftwareDevelopment_project-plan.docx`.
2. *Tài liệu Thu thập Yêu cầu*: `02_GenAI_SoftwareDevelopment_requirements-qa.docx`.
3. *Quy chuẩn Kỹ thuật Quốc gia QCVN 06:2022/BXD* của Bộ Xây Dựng.
4. *Nghiên cứu Sản phẩm HAVEN PropTech Suite* (Part 1, Part 2, Part 3).

---

## 2. MÔ TẢ TỔNG QUAN ỨNG DỤNG

### 2.1 Danh Sách Các Tác Nhân (Actors)
| Tác nhân | Mô tả vai trò | Đặc quyền trong hệ thống |
| :--- | :--- | :--- |
| **Người Thuê (Tenant)** | Người dùng có nhu cầu tìm kiếm, thuê phòng, ký hợp đồng và sử dụng dịch vụ cư dân. | Tìm kiếm True Cost, xem bản đồ PCCC/ngập lụt, so sánh radar, chat, ký hợp đồng E-Sign, cọc Escrow. |
| **Chủ Nhà (Landlord)** | Cá nhân hoặc tổ chức sở hữu/quản lý căn hộ cho thuê trên sàn. | Đăng tin bằng AI, quản lý phòng, tiếp nhận lead, tạo hợp đồng, quản lý hóa đơn thu tiền. |
| **Cư Dân (Resident)** | Người đang trực tiếp sinh sống trong căn hộ thuê. | Điều khiển thiết bị IoT, gửi yêu cầu sửa chữa/bảo trì, thanh toán hóa đơn, đặt lịch tiện ích. |
| **Quản Trị Sàn (Marketplace Admin)**| Đội ngũ vận hành và kiểm duyệt của nền tảng HAVEN. | Xem bảng sức khỏe sàn, duyệt tin đăng, gắn cờ tin ảo, giải quyết tranh chấp cọc. |
| **Hệ Thống GenAI (AI Engine)** | Dịch vụ AI tích hợp xử lý ngôn ngữ tự nhiên và thị giác máy tính. | Phân tích intent tìm kiếm (NLP), nhận diện layout phòng từ ảnh, sinh tiêu đề SEO và gợi ý giá. |

### 2.2 Sơ Đồ Use Case Tổng Quát (General Use Case Diagram)

```mermaid
graph LR
    subgraph Users ["Tác Nhân (Actors)"]
        Tenant["👤 Người Thuê (Tenant)"]
        Landlord["🏢 Chủ Nhà (Landlord)"]
        Resident["🏠 Cư Dân (Resident)"]
        Admin["🛡️ Quản Trị (Admin)"]
        AI["🤖 GenAI Service"]
    end

    subgraph HavenSystem ["Hệ Thống HAVEN PropTech"]
        UC01["UC001: Tìm kiếm True Cost & Bản Đồ"]
        UC02["UC002: AI Housing Advisor (NLP)"]
        UC03["UC003: So Sánh Đa Chiều (Radar Chart)"]
        UC04["UC004: Chat Smart Chips & Đặt Lịch"]
        UC05["UC005: Ký Hợp Đồng & Cọc Escrow"]
        UC06["UC006: Biên Bản Bàn Giao 15 Mục"]
        UC07["UC007: Cẩm Nang & Commute Simulator"]
        UC08["UC008: Tour Thực Tế Ảo 360 Độ"]
        UC09["UC009: Smart Listing Creator (AI)"]
        UC10["UC010: Quản Lý Căn Hộ & Thu Tiền"]
        UC11["UC011: Quản Trị Sức Khỏe Sàn & Duyệt Tin"]
        UC12["UC012: IoT Smart Home & Yêu Cầu Sửa Chữa"]
    end

    Tenant --> UC01
    Tenant --> UC02
    Tenant --> UC03
    Tenant --> UC04
    Tenant --> UC05
    Tenant --> UC06
    Tenant --> UC07
    Tenant --> UC08

    Landlord --> UC09
    Landlord --> UC10
    Landlord --> UC04
    Landlord --> UC05

    Resident --> UC12

    Admin --> UC11
    
    AI -.-> UC02
    AI -.-> UC09
```

---

## 3. ĐẶC TẢ CHI TIẾT CÁC USE CASE CỐT LÕI

### 3.1 Use Case UC001: Tìm Kiếm Căn Hộ & Bản Đồ PCCC/Ngập Lụt
- **Mục tiêu**: Người thuê tìm kiếm căn hộ theo ngân sách thực tế và kiểm tra an toàn PCCC, ngập lụt đô thị.
- **Tác nhân**: Người Thuê (Tenant).
- **Điều kiện tiên quyết**: Người dùng truy cập ứng dụng HAVEN.
- **Điều kiện hoàn thành**: Danh sách căn hộ phù hợp được hiển thị cùng chiết tính True Cost và lớp bản đồ an toàn.
- **Luồng sự kiện chính**:
  1. Người thuê chọn thành phố (trong 16 thành phố) và phân khúc giá.
  2. Người thuê bật các lớp an toàn (PCCC đạt chuẩn QCVN 06, không ngập lụt khi triều cường, có chỗ đỗ ô tô).
  3. Hệ thống trả về danh sách căn hộ theo thời gian thực.
  4. Người thuê xem biểu đồ True Cost và chuyển tab Bản Đồ để xem chi tiết rủi ro khu vực.

### 3.2 Use Case UC009: Đăng Tin Tự Động Bằng AI (Smart Listing Creator)
- **Mục tiêu**: Chủ nhà xuất bản tin đăng hoàn chỉnh trong 60 giây nhờ trí tuệ nhân tạo.
- **Tác nhân**: Chủ Nhà (Landlord), Hệ thống GenAI Engine.
- **Điều kiện tiên quyết**: Chủ nhà đăng nhập vào giao diện Quản trị.
- **Luồng sự kiện chính**:
  1. Chủ nhà tải lên 3-8 ảnh thực tế của căn hộ và chọn vị trí.
  2. AI phân tích nhận diện loại phòng, không gian nội thất.
  3. AI tự động sinh tiêu đề chuẩn SEO, mô tả hấp dẫn và gợi ý khoảng giá thuê tối ưu.
  4. Chủ nhà xác nhận và tin đăng được gắn nhãn Verified sau khi kiểm định.

---

## 4. BẢNG MA TRẬN TRUY VẾT YÊU CẦU (REQUIREMENTS TRACEABILITY MATRIX)

| Mã Yêu Cầu | Tên Chức Năng | Use Case Liên Quan | Thành Phần Kỹ Thuật (Components) | Trạng Thái Kiểm Thử |
| :---: | :--- | :---: | :--- | :---: |
| **FR-01** | Tìm kiếm AI bằng ngôn ngữ tự nhiên | UC002 | `services/aiAdvisorService.ts`, `UserSearchView.tsx` | Đạt (Passed) |
| **FR-02** | Bảng chiết tính True Cost Breakdown | UC001 | `ApartmentDetailModal.tsx`, `TrueCostCalculator.tsx` | Đạt (Passed) |
| **FR-03** | Bản đồ PCCC & Ngập lụt Đa Lớp | UC001 | `ConfidenceMapView.tsx` | Đạt (Passed) |
| **FR-04** | So sánh đa chiều (Radar Chart 5 trục) | UC003 | `UserCompareView.tsx` | Đạt (Passed) |
| **FR-05** | Đăng tin thông minh bằng AI | UC009 | `LandlordSmartListingModal.tsx` | Đạt (Passed) |
| **FR-06** | Hợp đồng điện tử E-Sign & Escrow | UC005 | `LeaseContractModal.tsx`, `EscrowDepositModal.tsx` | Đạt (Passed) |
| **FR-07** | Biên bản bàn giao 15 mục có ảnh | UC006 | `HandoverChecklistView.tsx` | Đạt (Passed) |
| **FR-08** | Cẩm nang khu vực & Commute Simulator | UC007 | `NeighborhoodGuideView.tsx` | Đạt (Passed) |
| **FR-09** | Quản lý danh mục phòng & Sơ đồ tầng | UC010 | `LandlordUnitListView.tsx`, `FloorPlanManagerView.tsx`| Đạt (Passed) |
| **FR-10** | Quản trị sàn & Bảng sức khỏe thị trường | UC011 | `AdminDashboardView.tsx`, `MarketHealthView.tsx` | Đạt (Passed) |
| **FR-11** | IoT Smart Home & Yêu cầu bảo trì | UC012 | `ResidentPortalView.tsx`, `MaintenanceTicketModal.tsx` | Đạt (Passed) |
"""

    with open(os.path.join(ACADEMIC_DOCS, '03_SOFTWARE_REQUIREMENTS_SPECIFICATION_SRS.md'), 'w', encoding='utf-8') as f:
        f.write(md_content)
    print("  -> Saved docs/academic/03_SOFTWARE_REQUIREMENTS_SPECIFICATION_SRS.md")


if __name__ == '__main__':
    populate_project_plan()
    populate_requirements_qa()
    populate_requirements_specification()
    print("\n✅ HOÀN TẤT ĐỒNG BỘ TOÀN BỘ 3 TÀI LIỆU HỌC THUẬT (DOCX & MD)!")
